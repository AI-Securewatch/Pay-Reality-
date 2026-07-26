"""Format-specific text extraction for the AI Policy Builder
(AI_EXTRACTION_PIPELINE.md Stage 2). Every supported format is normalized
to one plain-text blob with inline location markers, so the LLM analysis
stage (claude_provider.py) never needs to know what format the source
document was.

This module has no knowledge of RuntimePolicy, the extraction provider,
or the database; it only turns bytes into marked-up text, the same
narrow-responsibility discipline domain/runtime_policy/ already holds
itself to.
"""

import csv
import io

import openpyxl
from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_FORMATS = ("pdf", "docx", "xlsx", "csv", "text")

_EXTENSION_TO_FORMAT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".csv": "csv",
    ".txt": "text",
}

_CONTENT_TYPE_TO_FORMAT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/csv": "csv",
    "text/plain": "text",
}


class UnsupportedFormatError(Exception):
    pass


def detect_format(filename: str, content_type: str | None) -> str:
    """Filename extension first (more reliable than a browser-supplied
    content_type, which is frequently generic or wrong for these formats),
    falling back to content_type."""
    lower_name = (filename or "").lower()
    for ext, fmt in _EXTENSION_TO_FORMAT.items():
        if lower_name.endswith(ext):
            return fmt
    if content_type in _CONTENT_TYPE_TO_FORMAT:
        return _CONTENT_TYPE_TO_FORMAT[content_type]
    raise UnsupportedFormatError(f"unsupported_format: {filename!r} ({content_type!r})")


def _extract_pdf(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for i, page in enumerate(reader.pages):
        parts.append(f"--- page {i + 1} ---\n{page.extract_text() or ''}")
    return "\n\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = []
    n = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        n += 1
        parts.append(f"--- paragraph {n} ---\n{para.text}")
    for table in doc.tables:
        for row in table.rows:
            n += 1
            cells = [c.text for c in row.cells]
            parts.append(f"--- paragraph {n} ---\n{chr(9).join(cells)}")
    return "\n\n".join(parts)


def _extract_xlsx(raw: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if all(cell is None for cell in row):
                continue
            cells = ["" if cell is None else str(cell) for cell in row]
            parts.append(f"--- sheet '{sheet.title}', row {row_idx} ---\n{chr(9).join(cells)}")
    return "\n\n".join(parts)


def _extract_csv(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    parts = []
    for row_idx, row in enumerate(csv.reader(io.StringIO(text)), start=1):
        if not any(cell.strip() for cell in row):
            continue
        parts.append(f"--- row {row_idx} ---\n{chr(9).join(row)}")
    return "\n\n".join(parts)


def _extract_text(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    return f"--- document ---\n{text}"


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "csv": _extract_csv,
    "text": _extract_text,
}


def extract_text(format: str, raw: bytes) -> str:
    """Returns marked-up plain text. An empty or all-blank document
    produces an empty string, a valid outcome (AI_EXTRACTION_PIPELINE.md
    Stage 2), not an error; callers treat zero candidates from empty text
    as a normal, successfully extracted (zero-result) upload."""
    if format not in _EXTRACTORS:
        raise UnsupportedFormatError(f"unsupported_format: {format!r}")
    return _EXTRACTORS[format](raw)
